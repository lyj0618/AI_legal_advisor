from __future__ import annotations

import importlib
import re
from pathlib import Path

from app.config import settings

# 法条切片逻辑版本（/health 与分块结果校验用）
CHUNKING_VERSION = "legal-article-v4"

# 法条编号（第X条）
_ARTICLE_NUM = r"第[一二三四五六七八九十百千万零壹贰叁肆伍陆柒捌玖拾佰仟\d]+条"
_LEGAL_ARTICLE_DETECT = re.compile(_ARTICLE_NUM)
_ARTICLE_LINE = re.compile(rf"^(?P<num>{_ARTICLE_NUM})\s*$", re.MULTILINE)
_ARTICLE_LINE_INLINE = re.compile(
    rf"^(?P<num>{_ARTICLE_NUM})\s+[\u4e00-\u9fff]",
    re.MULTILINE,
)
_STRUCT_LINE = re.compile(
    r"^第[一二三四五六七八九十百千万零壹贰叁肆伍陆柒捌玖拾佰仟\d]+[编章节]\s*$"
)
_PUNCT_ONLY = re.compile(r"^[,，、.;；…—\-]+$")
_CN_NUM = r"[一二三四五六七八九十百千万\d]+"


def extract_text(file_path: Path, name: str) -> str:
    """仅返回正文（供预览等只需文本的调用方使用）。"""
    return extract_text_with_images(file_path, name, None)[0]


def extract_text_with_images(
    file_path: Path, name: str, doc_id: str | None
) -> tuple[str, list[str], list[list[str]]]:
    """返回 (正文文本, 文档内嵌图片文件名列表, 按章节归类的图片文件名列表)。

    doc_id 为 None 时不提取图片（预览场景）。docx 场景会按文档顺序把图片
    保存到 data_path/doc_images/{doc_id}/，并在返回的图片名列表中保持同一顺序。
    image_sections 按编号章节（如 1.1.2）把图片分组，便于分块时把截图和对应问题条目放在一起。
    """
    suffix = Path(name).suffix.lower()
    if suffix in (".txt", ".md", ".markdown", ".csv"):
        return file_path.read_text(encoding="utf-8", errors="ignore"), []
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            from app.services.pdf_ocr import extract_pdf_text_with_ocr_fallback

            reader = PdfReader(str(file_path))
            plain = "\n".join(page.extract_text() or "" for page in reader.pages)
            text, _source = extract_pdf_text_with_ocr_fallback(file_path, plain)
            return text, []
        except Exception as e:
            raise ValueError(f"PDF 解析失败: {e}") from e
    if suffix in (".docx", ".doc"):
        return _extract_word(file_path, suffix, doc_id=doc_id)
    raise ValueError(f"暂不支持该文件类型: {suffix}，支持 txt/md/pdf/docx")


def _para_list_level(para) -> int | None:
    """返回段落的列表级别 ilvl，没有列表编号时返回 None。"""
    pPr = para._p.pPr
    if pPr is None or pPr.numPr is None or pPr.numPr.ilvl is None:
        return None
    return pPr.numPr.ilvl.val


def _reconstruct_list_number(counters: dict[int, int], ilvl: int) -> str:
    """根据当前各级计数器重建 Word 自动编号（如 1.1.2）。"""
    counters[ilvl] = counters.get(ilvl, 0) + 1
    for k in list(counters.keys()):
        if k > ilvl:
            counters[k] = 0
    return ".".join(str(counters.get(l, 0)) for l in range(ilvl + 1))


def _extract_word(
    file_path: Path, suffix: str, doc_id: str | None = None
) -> tuple[str, list[str], list[list[str]]]:
    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(file_path))
            parts: list[str] = []
            list_counters: dict[int, int] = {}
            # 记录三级编号段落（问题条目）在原始段落序列中的位置，用于把图片归到对应条目
            section_starts: list[int] = []
            for para_idx, para in enumerate(doc.paragraphs):
                text = (para.text or "").strip()
                if not text:
                    continue
                ilvl = _para_list_level(para)
                if ilvl is not None:
                    # Word 自动编号不会出现在 para.text 中，按顺序重建编号前缀
                    text = f"{_reconstruct_list_number(list_counters, ilvl)} {text}"
                    if ilvl >= 2:
                        section_starts.append(para_idx)
                parts.append(text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [(cell.text or "").strip() for cell in row.cells]
                    line = " | ".join(c for c in cells if c)
                    if line:
                        parts.append(line)
            content = "\n".join(parts).strip()
            if not content:
                raise ValueError("Word 文档无有效文本内容")

            image_marks: list[str] = []
            image_sections: list[list[str]] = []
            if doc_id:
                image_positions = _extract_docx_image_positions(doc, doc_id)
                image_marks = [fn for fn, _ in image_positions]
                image_sections = [[] for _ in section_starts]
                for fn, pos in image_positions:
                    if pos < 0 or not section_starts:
                        # 无法定位或没有章节条目时，归到最后一个条目或独立列表
                        if image_sections:
                            image_sections[-1].append(fn)
                        continue
                    # 找到该图片所属的问题条目：最后一个不大于图片段落位置的 section_start
                    target = -1
                    for i, start in enumerate(section_starts):
                        if start <= pos:
                            target = i
                        else:
                            break
                    if target >= 0:
                        image_sections[target].append(fn)
                    else:
                        image_sections[0].append(fn)
            return content, image_marks, image_sections
        except ImportError as e:
            raise ValueError("未安装 python-docx，请执行 pip install python-docx") from e
        except Exception as e:
            raise ValueError(f"Word(.docx) 解析失败: {e}") from e

    raise ValueError("旧版 .doc 格式请另存为 .docx 后上传")


def _extract_docx_image_positions(doc, doc_id: str | None) -> list[tuple[str, int]]:
    """按文档顺序提取 docx 内嵌图片，返回 (文件名, 所在段落索引) 列表。

    段落索引用于把图片归到正确的章节/问题条目里；无法定位的（页眉/页脚/表格等）
    返回索引 -1。
    """
    if not doc_id:
        return []
    from docx.oxml.ns import qn

    out_dir = settings.data_path / "doc_images" / doc_id
    out_dir.mkdir(parents=True, exist_ok=True)
    marks: list[tuple[str, int]] = []
    seen: set[str] = set()

    # 建立段落元素 -> 段落索引的映射
    para_el_to_idx = {para._p: idx for idx, para in enumerate(doc.paragraphs)}

    def _para_idx_for_blip(blip) -> int:
        el = blip
        while el is not None:
            if el.tag.endswith("}p"):
                return para_el_to_idx.get(el, -1)
            el = el.getparent()
        return -1

    try:
        # 按 body XML 顺序扫描 <a:blip>，同时记录所属段落
        for blip in doc.element.body.findall(".//" + qn("a:blip")):
            r_embed = blip.get(qn("r:embed"))
            if not r_embed or r_embed in seen:
                continue
            seen.add(r_embed)
            rel = doc.part.rels.get(r_embed)
            if not rel or "image" not in str(rel.reltype):
                continue
            ext = Path(str(rel.target_ref)).suffix.lower() or ".png"
            fn = f"img_{len(marks):03d}{ext}"
            out_dir.joinpath(fn).write_bytes(rel.target_part.blob)
            marks.append((fn, _para_idx_for_blip(blip)))
    except Exception:
        # 兜底：遍历所有图片关系（顺序可能略有差异，但保证不丢图）
        for rid, rel in getattr(doc.part, "rels", {}).items():
            if "image" not in str(rel.reltype) or rid in seen:
                continue
            seen.add(rid)
            ext = Path(str(rel.target_ref)).suffix.lower() or ".png"
            fn = f"img_{len(marks):03d}{ext}"
            try:
                out_dir.joinpath(fn).write_bytes(rel.target_part.blob)
            except Exception:
                continue
            marks.append((fn, -1))
    return marks


def _extract_docx_images(doc, doc_id: str | None) -> list[str]:
    """仅返回文件名列表的兼容接口。"""
    return [fn for fn, _ in _extract_docx_image_positions(doc, doc_id)]


def collapse_vertical_legal_headers(text: str) -> str:
    """合并 PDF 竖排断行：第\\n一\\n条 → 第一条。"""
    if not text:
        return text
    lines = text.split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if s == "第" and i + 2 < len(lines):
            mid = lines[i + 1].strip()
            end = lines[i + 2].strip()
            if re.fullmatch(_CN_NUM, mid) and end in ("条", "编", "章", "节"):
                out.append(f"第{mid}{end}")
                i += 3
                continue
        out.append(s)
        i += 1
    merged = "\n".join(out)
    merged = re.sub(
        rf"第\s*({_CN_NUM})\s*条",
        r"第\1条",
        merged,
    )
    return merged


def _is_pdf_broken_layout(text: str) -> bool:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if len(lines) < 80:
        return False
    avg_len = sum(len(ln) for ln in lines) / len(lines)
    short = sum(1 for ln in lines if len(ln) <= 2)
    return avg_len < 22 or short > len(lines) * 0.25


def reflow_pdf_legal_text(text: str) -> str:
    """合并 PDF 逐行抽取的断行。"""
    if not text:
        return text
    text = collapse_vertical_legal_headers(text)
    if not _is_pdf_broken_layout(text):
        return text

    lines = text.split("\n")
    out: list[str] = []
    buf: list[str] = []

    def flush_buf() -> None:
        if buf:
            out.append("".join(buf))
            buf.clear()

    for raw in lines:
        s = raw.strip()
        if not s:
            flush_buf()
            if out and out[-1] != "":
                out.append("")
            continue
        s = re.sub(r"(?<=[\u4e00-\u9fff]) (?=[\u4e00-\u9fff])", "", s)
        if _PUNCT_ONLY.match(s):
            if buf:
                buf[-1] += s
            continue
        if _STRUCT_LINE.match(s) or _ARTICLE_LINE.match(s):
            flush_buf()
            out.append(s)
        elif _ARTICLE_LINE_INLINE.match(s):
            flush_buf()
            out.append(s)
        else:
            buf.append(s)
    flush_buf()
    return "\n".join(out).strip()


def _trim_toc_before_body(text: str) -> str:
    anchor = re.search(
        r"(?:^|\n)\s*第[一1]\s*编[\s\S]{0,800}?(?:^|\n)\s*第一条\s*(?:\n|$)",
        text,
        re.MULTILINE,
    )
    if anchor:
        return text[anchor.start() :].lstrip()
    anchor = re.search(
        r"(?:^|\n)\s*第一章[\s\S]{0,600}?(?:^|\n)\s*第一条\s*(?:\n|$)",
        text,
        re.MULTILINE,
    )
    if anchor:
        return text[anchor.start() :].lstrip()

    for m in _ARTICLE_LINE.finditer(text):
        tail = text[m.end() : m.end() + 400]
        plain = re.sub(r"[\s,，。；、…—\-]", "", tail)
        if len(plain) >= 30:
            return text[m.start() :].lstrip()
    return text


def prepare_legal_text(text: str) -> str:
    """清洗/分块/预览统一使用的正文整理（不依赖重启后端）。"""
    text = reflow_pdf_legal_text(text or "")
    return _trim_toc_before_body(text)


def _is_legal_document(text: str) -> bool:
    return len(_LEGAL_ARTICLE_DETECT.findall(prepare_legal_text(text))) >= 2


# 编号章节切片：Word 自动编号（如 1.1.2 新建项目...）
# 注意：编号已在 _extract_word 中按段落顺序重建；这里匹配至少 3 级编号
_NUMBERED_HEADING = re.compile(r"^\s*(?:\d+\.){2,}\d+\s+\S", re.MULTILINE)


def _has_numbered_sections(text: str) -> bool:
    return len(_NUMBERED_HEADING.findall(text or "")) >= 2


def _split_numbered_sections(text: str) -> list[str]:
    text = text or ""
    matches = list(_NUMBERED_HEADING.finditer(text))
    if len(matches) < 2:
        return []
    segments: list[str] = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        seg = text[start:end].strip()
        if seg:
            segments.append(seg)
    return segments


def _collect_article_starts(text: str) -> list[int]:
    starts: list[int] = []
    for m in _ARTICLE_LINE.finditer(text):
        starts.append(m.start())
    for m in _ARTICLE_LINE_INLINE.finditer(text):
        if m.start() not in starts:
            starts.append(m.start())
    starts.sort()
    deduped: list[int] = []
    for p in starts:
        if not deduped or p - deduped[-1] > 3:
            deduped.append(p)
    return deduped


def _split_legal_articles(text: str) -> list[str]:
    text = prepare_legal_text(text)
    starts = _collect_article_starts(text)
    if len(starts) < 2:
        parts = [s.strip() for s in re.split(rf"(?={_ARTICLE_NUM})", text) if s.strip()]
        parts = [p for p in parts if re.match(rf"^{_ARTICLE_NUM}", p)]
        if len(parts) >= 2:
            return parts
        return []

    segments: list[str] = []
    preamble = text[: starts[0]].strip()
    if preamble and len(preamble) > 20:
        segments.append(preamble)

    for i, start in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(text)
        block = text[start:end].strip()
        if block:
            segments.append(block)
    return segments


def validate_legal_chunk_parts(parts: list[str], prepared_text: str) -> bool:
    """校验法条文档是否已按条切片（避免旧逻辑大块写入）。"""
    if not parts:
        return False
    expected = len(_LEGAL_ARTICLE_DETECT.findall(prepared_text))
    article_chunks = sum(
        1 for p in parts if re.match(rf"^{_ARTICLE_NUM}", (p or "").strip())
    )
    max_len = max(len(p) for p in parts)
    if expected >= 20:
        return article_chunks >= max(20, int(expected * 0.5)) and max_len < 800
    return article_chunks >= 2 and max_len < 1200


def _split_oversized(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]

    parts: list[str] = []
    for para in re.split(r"\n{2,}", text):
        para = para.strip()
        if not para:
            continue
        if len(para) <= max_chars:
            parts.append(para)
        else:
            for i in range(0, len(para), max_chars):
                parts.append(para[i : i + max_chars])
    return parts or [text[:max_chars]]


def _merge_segments(segments: list[str], max_chars: int) -> list[str]:
    parts: list[str] = []
    buf = ""
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        if len(buf) + len(seg) + 1 <= max_chars:
            buf = f"{buf}\n{seg}".strip() if buf else seg
        else:
            if buf:
                parts.append(buf)
            if len(seg) > max_chars:
                parts.extend(_split_oversized(seg, max_chars))
                buf = ""
            else:
                buf = seg
    if buf:
        parts.append(buf)
    return parts


def split_chunks(
    text: str,
    chunk_token_num: int = 512,
    delimiter: str = "",
    *,
    chunk_strategy: str = "auto",
) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []

    max_chars = max(chunk_token_num * 2, 200)
    use_legal = chunk_strategy == "legal_article" or (
        chunk_strategy not in ("naive", "numbered_section") and _is_legal_document(text)
    )

    if delimiter and chunk_strategy != "legal_article":
        seps = [s.strip() for s in delimiter.split("\\n") if s.strip()]
        if not seps:
            seps = [delimiter]
        pattern = "|".join(re.escape(s) for s in seps)
        segments = [s.strip() for s in re.split(pattern, text) if s.strip()]
        parts = _merge_segments(segments, max_chars)
    elif chunk_strategy == "numbered_section" or (
        chunk_strategy not in ("legal_article", "naive") and _has_numbered_sections(text)
    ):
        segments = _split_numbered_sections(text)
        parts = []
        for seg in segments:
            parts.extend(_split_oversized(seg, max_chars))
    elif use_legal:
        prepared = prepare_legal_text(text)
        segments = _split_legal_articles(prepared)
        if not segments:
            segments = _split_legal_articles(text)
        parts = []
        for seg in segments:
            parts.extend(_split_oversized(seg, max_chars))
        if not validate_legal_chunk_parts(parts, prepared):
            prepared2 = prepare_legal_text(collapse_vertical_legal_headers(text))
            segments2 = _split_legal_articles(prepared2)
            parts2 = []
            for seg in segments2:
                parts2.extend(_split_oversized(seg, max_chars))
            if validate_legal_chunk_parts(parts2, prepared2):
                parts = parts2
    else:
        segments = [s.strip() for s in re.split(r"\n{2,}", text) if s.strip()]
        parts = _merge_segments(segments, max_chars)

    if not parts:
        for i in range(0, len(text), max_chars):
            parts.append(text[i : i + max_chars])
    return parts


def reload_chunking_module():
    """每次清洗/分块前加载最新切片代码，无需重启 uvicorn。"""
    import app.services.chunking as mod

    return importlib.reload(mod)
